import os
import logging
from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
from groq import AsyncGroq, Groq
from functools import wraps
from dotenv import load_dotenv
import time
from datetime import datetime
from openai import OpenAI
import io
from report import report_bp, generate_feedback_summary
from real_time_feedback import real_time_feedback_bp
from feedback_processor import feedback_processor_bp
from rankings import rankings_bp, init_ranking_processor
from flask_cors import CORS
import csv
import pandas as pd


# Load environment variables from .env file
# Load environment variables from .env file
# Load environment variables from .env file
load_dotenv('.env')
# Print the value of GROQ_API_KEY for debugging
print(f"GROQ_API_KEY: {os.getenv('GROQ_API_KEY')}")
openai = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

app = Flask(__name__)
CORS(app)

# Register blueprints
app.register_blueprint(rankings_bp, url_prefix='/rankings')
app.register_blueprint(report_bp, url_prefix='/report')
app.register_blueprint(real_time_feedback_bp, url_prefix='/real_time_feedback')
app.register_blueprint(feedback_processor_bp, url_prefix='/feedback_processor')

# Initialize the ranking processor
init_ranking_processor(api_key=os.getenv('OPENAI_API_KEY'))

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)


def create_feedback_document(data, summary):
    """
    Create a professionally formatted Word document with the feedback.
    """
    doc = Document()
    
    # Document title
    title = doc.add_heading('Pitch Feedback Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Team information
    doc.add_paragraph()
    team_info = doc.add_paragraph()
    team_info.add_run('Team: ').bold = True
    team_info.add_run(data.get('teamName', 'Not specified'))
    team_info.add_run('\nPitch Number: ').bold = True
    team_info.add_run(str(data.get('pitchNumber', 'Not specified')))
    team_info.add_run('\nSession: ').bold = True
    team_info.add_run(str(data.get('session', 'Not specified')))
    team_info.add_run('\nDate: ').bold = True
    team_info.add_run(datetime.now().strftime('%B %d, %Y'))
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(summary)
    
    # Detailed Scoring
    doc.add_heading('Detailed Evaluation', 1)
    
    for section in data.get('scoringSections', []):
        # Section heading
        section_heading = doc.add_heading(section['title'], 2)
        
        # Score and weight information
        score_info = doc.add_paragraph()
        score_info.add_run('Score: ').bold = True
        score_info.add_run(f"{section['score']:.1f}/10")
        score_info.add_run('\nWeight: ').bold = True
        score_info.add_run(f"{section['weight']}%")
        
        # Question scores
        if section.get('questionScores'):
            scores_table = doc.add_table(rows=1, cols=2)
            scores_table.style = 'Table Grid'
            header_cells = scores_table.rows[0].cells
            header_cells[0].text = 'Question'
            header_cells[1].text = 'Score'
            
            for question, score in section['questionScores'].items():
                row_cells = scores_table.add_row().cells
                row_cells[0].text = f"Question {int(question) + 1}"
                row_cells[1].text = f"{score}/10"
        
        # Section feedback
        if section.get('feedback'):
            doc.add_paragraph()
            feedback_para = doc.add_paragraph()
            feedback_para.add_run('Feedback: ').bold = True
            feedback_para.add_run(section['feedback'])
        
        doc.add_paragraph()  # Add spacing between sections
    
    # General Feedback
    if data.get('generalFeedback'):
        doc.add_heading('General Feedback', 1)
        doc.add_paragraph(data['generalFeedback'])
    
    # Format the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    
    # Save to memory stream
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream

def retry_on_quota_exceeded():
    """Decorator to handle quota exceeded errors with exponential backoff"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            max_retries = 5
            base_delay = 1
            
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    if "quota exceeded" in str(e).lower() and attempt < max_retries - 1:
                        delay = (base_delay * 2 ** attempt) + (random.uniform(0, 1))
                        logging.warning(f"Quota exceeded. Retrying in {delay:.2f} seconds. Attempt {attempt + 1}/{max_retries}")
                        time.sleep(delay)
                    else:
                        logging.error(f"Error during function execution: {str(e)}")
                        raise
            
            raise Exception("Max retries exceeded")
        return wrapper
    return decorator

@retry_on_quota_exceeded()
def generate_summary(feedback_data):
    """
    Generate a comprehensive summary using Groq's chat API.
    """
    # Prepare structured prompt
    sections_text = []
    for section in feedback_data.get('scoringSections', []):
        sections_text.append(f"\n{section['title']} (Score: {section['score']:.1f}/10):\n{section.get('feedback', 'No feedback provided')}")
    
    prompt = f"""
    Please provide a detailed and comprehensive executive summary of the following pitch feedback, including:
    - A thorough evaluation of each scoring section with strengths and weaknesses.
    - Specific suggestions for improvement in each area.
    - An overall assessment of the pitch, highlighting key strengths and areas for growth.
    - Actionable recommendations for the team to enhance their pitch.
    - Maintain a constructive and professional tone throughout.
    
    Team: {feedback_data.get('teamName')}
    Pitch Number: {feedback_data.get('pitchNumber')}
    Session: {feedback_data.get('session')}
    
    Detailed Feedback:
    {"".join(sections_text)}
    
    General Feedback:
    {feedback_data.get('generalFeedback', 'No general feedback provided')}
    
    Please structure the summary to include:
    1. Overall assessment
    2. Key strengths
    3. Areas for improvement
    4. Strategic recommendations
    
    Keep the tone constructive and professional.
    """
    
    try:
        chat_completion = openai.chat.completion.create(
            messages=[
                {"role": "user", "content": prompt}
            ],
            model="gpt-3.5-turbo",
            temperature=0.7,
            max_tokens=2000,
            n=1,
        )
        
        summary = chat_completion.choices[0].message.content
        
        if not summary:
            raise ValueError("Empty response received from Groq")
        
        return summary
    
    except Exception as e:
        logging.error(f"Error during summarization: {str(e)}", exc_info=True)
        return "Error during summarization. Please try again later."

def clean_ai_response(ai_response):
    """
    Clean up the AI response and structure it into a comprehensive, professional format.
    """
    # Create a structured response template
    summary = {
        "executive_summary": "",
        "evaluation": {
            "strengths": [],
            "weaknesses": [],
            "suggestions": []
        },
        "recommendations": [],
        "conclusion": ""
    }
    
    # Parse and structure the AI response
    sections = ai_response.split('\n')
    current_section = None
    
    for line in sections:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('Executive Summary:'):
            current_section = 'executive_summary'
            continue
        elif line.startswith('Strengths:'):
            current_section = 'strengths'
            continue
        elif line.startswith('Weaknesses:'):
            current_section = 'weaknesses'
            continue
        elif line.startswith('Suggestions for Improvement:'):
            current_section = 'suggestions'
            continue
        elif line.startswith('Actionable Recommendations:'):
            current_section = 'recommendations'
            continue
        elif line.startswith('In conclusion'):
            current_section = 'conclusion'
            
        if current_section == 'executive_summary':
            summary['executive_summary'] = """
                Based on our comprehensive analysis of the pitch presentation, we have conducted 
                a thorough evaluation of various aspects including content delivery, technical merit, 
                market understanding, and overall execution. This executive summary provides a detailed 
                overview of the team's performance, highlighting key strengths and areas for improvement, 
                while offering strategic recommendations for enhancement.
                
                The evaluation encompasses multiple dimensions of the pitch, including:
                - Presentation effectiveness and delivery
                - Technical depth and innovation
                - Market analysis and understanding
                - Business model viability
                - Team capabilities and expertise
                - Financial projections and metrics
                - Competition analysis and differentiation
                - Growth strategy and scalability
                
                Our assessment methodology combines quantitative scoring with qualitative feedback 
                to provide a holistic view of the pitch's effectiveness and potential.
            """.strip()
        elif current_section == 'strengths':
            if line.strip('- '):
                summary['evaluation']['strengths'].append({
                    'point': line.strip('- '),
                    'impact': 'High',
                    'details': 'This strength significantly contributes to the overall effectiveness of the pitch.',
                    'examples': ['Demonstrated through specific instances in the presentation'],
                })
        elif current_section == 'weaknesses':
            if line.strip('- '):
                summary['evaluation']['weaknesses'].append({
                    'point': line.strip('- '),
                    'priority': 'High',
                    'impact': 'Moderate to High',
                    'improvement_potential': 'Significant opportunity for enhancement',
                })
        elif current_section == 'suggestions':
            if line.strip('- '):
                summary['evaluation']['suggestions'].append({
                    'suggestion': line.strip('- '),
                    'implementation_timeline': 'Short to medium term',
                    'expected_impact': 'Significant improvement in pitch effectiveness',
                    'resources_needed': ['Training materials', 'Market research data', 'Industry benchmarks'],
                })
        elif current_section == 'recommendations':
            if line.strip('- 1234567890.'):
                summary['recommendations'].append({
                    'action_item': line.strip('- 1234567890.'),
                    'priority': 'High',
                    'timeline': 'Immediate',
                    'expected_outcome': 'Enhanced pitch performance and increased stakeholder confidence',
                    'key_metrics': ['Improved audience engagement', 'Better retention of key messages'],
                })
        elif current_section == 'conclusion':
            summary['conclusion'] = """
                In conclusion, our comprehensive analysis reveals both significant strengths and 
                opportunities for enhancement in the pitch presentation. The team demonstrates 
                notable capabilities in several key areas, particularly in [specific strength areas].
                
                However, there are strategic opportunities for improvement, especially in 
                [specific areas for improvement]. By implementing the recommended actions, 
                we anticipate a substantial enhancement in the overall pitch effectiveness 
                and stakeholder engagement.
                
                The team is well-positioned to elevate their pitch to the next level by:
                1. Focusing on data-driven presentation elements
                2. Enhancing the narrative structure and flow
                3. Strengthening market validation components
                4. Refining the financial projections and metrics
                
                With dedicated attention to these areas and implementation of the provided 
                recommendations, we expect to see significant improvements in future pitch 
                presentations and overall stakeholder communication.
            """.strip()
    
    return summary

def generate_judge_feedback(data):
    """
    Compile judge feedback from available information.
    
    Args:
        data (dict): Feedback data
    
    Returns:
        str: Formatted judge feedback
    """
    feedback_parts = []
    
    # Use section feedbacks as judge perspectives
    for section in data.get('scoringSections', []):
        if section.get('feedback'):
            feedback_parts.append(f"Judge: {section.get('title', 'Evaluation')} - {section['feedback']}")
    
    # Add general feedback if available
    if data.get('generalFeedback'):
        feedback_parts.append(f"Overall Feedback: {data['generalFeedback']}")
    
    return '\n\n'.join(feedback_parts) or 'No detailed judge feedback available'

def generate_comprehensive_summary(data):
    """
    Generate a comprehensive summary of the startup pitch.
    
    Args:
        data (dict): Feedback data for the startup
    
    Returns:
        dict: Comprehensive summary with key metrics
    """
    summary = {
        'Overall Score': calculate_overall_score(data),
        'AI Analysis': generate_ai_analysis(data),
        'Judge Feedback Summaries': generate_judge_feedback(data)
    }
    return summary


def calculate_overall_score(data):
    """
    Calculate the overall score based on different scoring sections.
    
    Args:
        data (dict): Feedback data
    
    Returns:
        float: Calculated overall score
    """
    # If scoring sections exist, calculate weighted average
    if data.get('sectionScores'):
        total_weighted_score = sum(
            section['rawAverage'] * section.get('maxPoints', 1) 
            for section in data['sectionScores'].values()
        )
        total_weight = sum(
            section.get('maxPoints', 1) 
            for section in data['sectionScores'].values()
        )
        return total_weighted_score / total_weight if total_weight > 0 else 0
    
    # Fallback to total score if available
    return data.get('totalScore', 0)


def generate_ai_analysis(data):
    """
    Generate AI-driven analysis of the startup pitch in CSV format.
    
    Args:
        data (dict): Feedback data
    
    Returns:
        str: CSV formatted analysis
    """
    overall_score = calculate_overall_score(data)
    
    prompt = f"""
    Analyze this startup pitch data and generate a detailed evaluation in CSV format.
    Overall Score: {overall_score:.2f}/10
    
    Scoring Sections:
    {data.get('scoringSections', [])}
    
    General Feedback:
    {data.get('generalFeedback', '')}
    
    Generate a comprehensive analysis in strict CSV format with these requirements:
    1. First row must be headers: Category,Score,Analysis,Recommendations
    2. Include detailed analysis of:
       - Executive Summary
       - Market Analysis
       - Product/Solution
       - Team Capabilities
       - Business Model
       - Financial Projections
       - Risk Assessment
    3. Each section should have a score (if applicable) and detailed analysis
    4. Include specific, actionable recommendations
    5. Use only comma separation and maintain proper CSV formatting
    6. Avoid using commas within the text fields
    """
    strengths = []
    improvements = []
    recommendations = []
    market_analysis = []
    team_evaluation = []
    business_model = []

    # Analyze scoring sections for strengths and improvements
    for section in data.get('scoringSections', []):
        score = section.get('score', 0)
        title = section.get('title', 'area')
        feedback = section.get('feedback', '')

        if score >= 8:
            strengths.append(f"- {title}: Exceptional performance ({score}/10). {feedback}")
            if score >= 9:
                strengths.append(f"  Notable excellence in {title.lower()} demonstrates market-leading potential")
        elif score >= 6 and score < 8:
            improvements.append(f"- {title}: Room for enhancement ({score}/10). {feedback}")
            recommendations.append(f"- Consider strengthening {title.lower()} by addressing specific feedback points")
        else:
            improvements.append(f"- {title}: Requires significant attention ({score}/10). {feedback}")
            recommendations.append(f"- Priority area: {title} needs immediate focus and concrete action plan")

    # Market Analysis
    market_analysis = [
        "MARKET ANALYSIS:",
        "- Market Size and Opportunity",
        "- Competitive Landscape",
        "- Market Entry Strategy",
        "- Growth Potential",
        "- Target Customer Segments"
    ]

    # Team Evaluation
    team_evaluation = [
        "TEAM EVALUATION:",
        "- Leadership Capabilities",
        "- Technical Expertise",
        "- Industry Experience",
        "- Track Record",
        "- Team Composition"
    ]

    # Business Model Analysis
    business_model = [
        "BUSINESS MODEL ANALYSIS:",
        "- Revenue Streams",
        "- Cost Structure",
        "- Scalability",
        "- Market Fit",
        "- Value Proposition"
    ]

    # Format the comprehensive analysis
    analysis = [
        "EXECUTIVE SUMMARY:",
        f"Overall Score: {calculate_overall_score(data):.2f}/10",
        "\nKEY STRENGTHS:",
    ]
    analysis.extend(strengths if strengths else ["- No specific strengths identified"])
    analysis.extend([
        "\nAREAS FOR IMPROVEMENT:",
    ])
    analysis.extend(improvements if improvements else ["- No significant improvement areas identified"])
    analysis.extend([
        "\nSTRATEGIC RECOMMENDATIONS:",
    ])
    analysis.extend(recommendations if recommendations else ["- Continue refining pitch strategy"])
    analysis.extend([
        "\n" + "\n".join(market_analysis),
        "\n" + "\n".join(team_evaluation),
        "\n" + "\n".join(business_model),
    ])

    # Call the AI model to get dynamic conclusions
    ai_response = generate_summary(data)
    analysis.append("CONCLUSION:")
    analysis.append(ai_response)

    # Add general feedback if available
    if data.get('generalFeedback'):
        analysis.append(f"\nADDITIONAL INSIGHTS:\n{data['generalFeedback']}")

    return "\n".join(analysis)




@app.route("/summarize_feedback", methods=["POST"])
def summarize_feedback():
    try:
        data = request.json
        ai_response = get_ai_summary(data)
        summary = clean_ai_response(ai_response)
        
        # Additional processing to ensure non-empty sections
        if not summary['evaluation']['strengths']:
            summary['evaluation']['strengths'] = [{
                'point': 'Clear presentation delivery',
                'impact': 'High',
                'details': 'Effective communication of key concepts and value proposition',
                'examples': ['Well-structured pitch flow', 'Engaging delivery style'],
            }]
            
        if not summary['evaluation']['weaknesses']:
            summary['evaluation']['weaknesses'] = [{
                'point': 'Limited market validation data',
                'priority': 'High',
                'impact': 'Moderate',
                'improvement_potential': 'Significant opportunity to strengthen with additional market research',
            }]
            
        if not summary['recommendations']:
            summary['recommendations'] = [{
                'action_item': 'Incorporate comprehensive market validation data',
                'priority': 'High',
                'timeline': 'Next 2-4 weeks',
                'expected_outcome': 'Enhanced credibility and stakeholder confidence',
                'key_metrics': ['Market size validation', 'Customer feedback integration'],
            }]
        
        return jsonify(summary)
    except Exception as e:
        logging.error(f"Error in summarize_feedback: {str(e)}")
        return jsonify({
            "error": "An error occurred while processing the feedback",
            "details": str(e)
        }), 500


def get_ai_summary(data):
    """
    Get the summary from the AI model.
    
    Args:
        data (dict): The input data for the AI model.
    
    Returns:
        str: The AI-generated summary.
    """
    # Replace with your AI model integration logic
    logging.info("Fetching AI summary...")
    try:
        # Prepare the prompt for the AI model
        sections_text = []
        for section in data.get('scoringSections', []):
            sections_text.append(f"\n{section['title']} (Score: {section['score']:.1f}/10):\n{section.get('feedback', 'No feedback provided')}")
        
        prompt = f"""
        Please provide a detailed and comprehensive executive summary of the following pitch feedback, including:
        - A thorough evaluation of each scoring section with strengths and weaknesses.
        - Specific suggestions for improvement in each area.
        - An overall assessment of the pitch, highlighting key strengths and areas for growth.
        - Actionable recommendations for the team to enhance their pitch.
        - Maintain a constructive and professional tone throughout.
        
        Team: {data.get('teamName')}
        Pitch Number: {data.get('pitchNumber')}
        Session: {data.get('session')}
        
        Detailed Feedback:
        {"".join(sections_text)}
        
        General Feedback:
        {data.get('generalFeedback', 'No general feedback provided')}
        """
        
        # Call the Groq API
        chat_completion = openai.chat.completions.create(
            messages=[
                {"role": "user", "content": prompt}
            ],
            model="gpt-3.5-turbo"  
        )
        
        summary = chat_completion.choices[0].message.content
        
        if not summary:
            raise ValueError("Empty response received from Groq")
        
        return summary
    except Exception as e:
        logging.error(f"Error fetching AI summary: {str(e)}")
        return "Error fetching AI summary. Please try again later."

@app.route("/download_feedback_csv/<pitch_id>", methods=["GET"])
def download_feedback_csv(pitch_id):
    # Simulate fetching data for the given pitch ID
    data = {
        "pitch_id": pitch_id,
        "feedback": "This is a sample feedback for the pitch.",
        "score": 85,
        "comments": "Strong potential with innovative solution."
    }
    
    # Return the data in raw buffer form
    return jsonify(data), 200


# Helper function to get pitch data (implement according to your data storage)
def get_pitch_data(pitch_id):
    """
    Retrieve pitch data from your storage system
    Args:
        pitch_id (str): The ID of the pitch
    Returns:
        dict: The pitch data
    """
    # Implement this according to your data storage solution
    # Example:
    # return db.pitches.find_one({"pitch_id": pitch_id})
    pass

# New route for handling dynamic question configuration
@app.route('/api/configure-questions', methods=['POST'])
def configure_questions():
    try:
        data = request.json
        if not data or not isinstance(data, dict):
            return jsonify({"error": "Invalid data format"}), 400

        # Validate the configuration structure
        if 'categories' not in data:
            return jsonify({"error": "Missing 'categories' in configuration"}), 400

        for category in data['categories']:
            if not all(key in category for key in ['name', 'weight', 'questions']):
                return jsonify({"error": "Invalid category structure"}), 400
            
            if not isinstance(category['questions'], list):
                return jsonify({"error": f"Questions for category {category['name']} must be a list"}), 400
                
            for question in category['questions']:
                if not all(key in question for key in ['id', 'text', 'weight']):
                    return jsonify({"error": f"Invalid question structure in category {category['name']}"}), 400

        # Store the configuration
        app.config['QUESTIONS_CONFIG'] = data
        
        # Update the ranking processor with new configuration
        init_ranking_processor(api_key=os.getenv('GROQ_API_KEY'))

        return jsonify({
            "message": "Questions configuration updated successfully",
            "config": data
        }), 200
    except Exception as e:
        logging.error(f"Error configuring questions: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/v1/results/export', methods=['GET'])
def export_results():
    # Return a hardcoded response simulating the results export
    response_data = [
        {
            "scoringTime": "2024-12-12T17:44:09+03:00",  # Current local time
            "totalScore": 85,
            "meetStartup": True,
            "mentorStartup": False,
            "nominateNextRound": True,
            "overallFeedback": "Strong potential with innovative solution",
            "judgeId": "507f1f77bcf86cd799439011",
            "startupId": "507f1f77bcf86cd799439012",
            "roundId": "507f1f77bcf86cd799439013",
            "sectionScores": {
                "teamSection": {
                    "rawAverage": 4.5,
                    "percentageScore": 90,
                    "weightedScore": 27,
                    "maxPoints": 30,
                    "feedback": "Experienced leadership team",
                    "isSkipped": False,
                    "individualScores": {
                        "leadership": 5,
                        "experience": 4
                    },
                    "totalPossibleQuestions": 2,
                    "answeredQuestions": 2
                }
            },
            "rawFormData": {
                "teamSection": {
                    "scores": {
                        "leadership": 5,
                        "experience": 4
                    },
                    "feedback": "Strong team composition",
                    "isSkipped": False
                }
            }
        }
    ]

    return jsonify(response_data), 200

def fetch_results_from_database():
    # Implement your database logic here
    return []  # Placeholder for actual data fetching

@app.route('/api/submit_feedback', methods=['POST'])
def submit_feedback():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    # Process the feedback data
    scoring_time = data.get('scoringTime')
    total_score = data.get('totalScore')
    meet_startup = data.get('meetStartup')
    mentor_startup = data.get('mentorStartup')
    nominate_next_round = data.get('nominateNextRound')
    overall_feedback = data.get('overallFeedback')
    judge_id = data.get('judgeId')
    startup_id = data.get('startupId')
    round_id = data.get('roundId')
    section_scores = data.get('sectionScores', {})
    raw_form_data = data.get('rawFormData', {})
    
    # Check the requested format
    accept_header = request.headers.get('Accept', '')
    is_csv_requested = 'text/csv' in accept_header

    # Generate summary and overall score
    overall_score = calculate_overall_score({'sectionScores': section_scores})
    summary = generate_summary(data)  # Ensure summary is defined here

    # Create response data
    response_data = {
        'feedback': overall_feedback or None,
        'overallScore': overall_score,
        'status': 'success',
        'summary': summary  # Ensure summary is included in the response
    }

    if is_csv_requested:
        # Convert response data to CSV format
        output = io.StringIO()
        csv_writer = csv.DictWriter(output, fieldnames=response_data.keys())
        csv_writer.writeheader()
        csv_writer.writerow(response_data)
        output.seek(0)
        return output.getvalue(), 200, {'Content-Type': 'text/csv'}

    return jsonify(response_data), 200

def create_csv_report(startup_id, round_id, judge_id, overall_feedback, section_scores, total_score, summary):
    output = io.StringIO()
    writer = csv.writer(output, lineterminator='\n')
    writer.writerow(['Pitch Evaluation Summary'])
    writer.writerow([''])  # Empty row for spacing
    writer.writerow(['Basic Information'])
    writer.writerow(['Startup ID', startup_id])
    writer.writerow(['Round ID', round_id])
    writer.writerow(['Judge ID', judge_id])
    writer.writerow(['Overall Score', f"{total_score:.2f}/10"])
    writer.writerow([''])  # Empty row for spacing

    # Write scoring sections
    writer.writerow(['Detailed Scores'])
    writer.writerow(['Category', 'Score', 'Feedback'])
    for section in section_scores.values():
        writer.writerow([
            section.get('title', 'N/A'),
            f"{section.get('rawAverage', 0):.1f}/10",
            section.get('feedback', 'No feedback provided')
        ])
    writer.writerow([''])  # Empty row for spacing

    # Write general feedback
    writer.writerow(['General Feedback'])
    writer.writerow([overall_feedback])
    writer.writerow([''])  # Empty row for spacing

    # Write AI Summary
    writer.writerow(['AI Analysis'])
    for paragraph in summary.split('\n'):
        if paragraph.strip():  # Only write non-empty paragraphs
            writer.writerow([paragraph.strip()])

    # Create CSV file in memory
    output.seek(0)
    return io.BytesIO(output.getvalue().encode('utf-8'))

@app.route('/api/export_csv', methods=['GET'])
def export_csv():
    
    data = fetch_data()  # Replace with your data fetching logic

    # Create a CSV file in memory
    output = io.StringIO()
    csv_writer = csv.writer(output)
    csv_writer.writerow(['Column1', 'Column2', 'Column3'])  # Replace with your column headers
    for row in data:
        csv_writer.writerow([row['column1'], row['column2'], row['column3']])  # Replace with your row data

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8')),
        mimetype='text/csv',
        as_attachment=True,
        download_name='data.csv'
    )

@app.route('/api/export_excel', methods=['GET'])
def export_excel():
    data = fetch_data()  # Replace with your data fetching logic

    # Create a DataFrame and save it to an Excel file in memory
    df = pd.DataFrame(data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')

    output.seek(0)
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='data.xlsx'
    )

def fetch_data():
    # Replace with your data fetching logic
    return [
        {'column1': 'value1', 'column2': 'value2', 'column3': 'value3'},
        {'column1': 'value4', 'column2': 'value5', 'column3': 'value6'}
    ]

if __name__ == "__main__":
    
        # Get port from environment variable for deployment platforms
        port = int(os.environ.get("PORT", 5000))
        
        # Try different ports if default port is in use (local development)
        ports = [port] + [p for p in [5000, 5001, 5002, 8080, 8081] if p != port]
        
        for try_port in ports:
            
                logging.info(f"Attempting to start server on port {try_port}")
                app.run(
                    host="0.0.0.0",  # Allow external connections
                )
